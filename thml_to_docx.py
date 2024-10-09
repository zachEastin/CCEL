import xml.etree.ElementTree as ET
from docx import Document
from docx.shared import Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.style import WD_STYLE_TYPE
from docx.opc.constants import RELATIONSHIP_TYPE


def add_hyperlink(paragraph, url, text):
    """
    Add a hyperlink to a paragraph.

    :param paragraph: The paragraph we are adding the hyperlink to.
    :param url: A string containing the required URL
    :param text: The text displayed for the link
    """
    # Create the w:hyperlink tag and add attributes
    part = paragraph.part
    r_id = part.relate_to(url, RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    # Create a run for the hyperlink text with formatting
    new_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")

    # Style the hyperlink (blue and underlined)
    c = OxmlElement("w:color")
    c.set(qn("w:val"), "0000FF")  # Blue color
    rPr.append(c)

    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")  # Single underline
    rPr.append(u)

    new_run.append(rPr)

    # Add the hyperlink text
    text_elem = OxmlElement("w:t")
    text_elem.text = text
    new_run.append(text_elem)

    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def add_footnote(paragraph, text):
    """
    Adds a footnote to the paragraph.
    """
    # Add a footnote reference to the paragraph
    footnote_reference = paragraph.add_run()
    footnote_reference._element.addfootnoteReference()

    # Access the footnotes part of the document
    part = paragraph.part
    footnotes_part = part.footnotes_part

    # Create a new footnote ID
    footnote_id = len(footnotes_part.element.findall(qn("w:footnote"))) + 1

    # Create the footnote element
    footnote = OxmlElement("w:footnote")
    footnote.set(qn("w:id"), str(footnote_id))

    # Add the footnote text
    footnote_p = OxmlElement("w:p")
    footnote_r = OxmlElement("w:r")
    footnote_t = OxmlElement("w:t")
    footnote_t.text = text
    footnote_r.append(footnote_t)
    footnote_p.append(footnote_r)
    footnote.append(footnote_p)

    # Append the footnote to the footnotes part
    footnotes_part.element.append(footnote)


def get_style_mapping(document):
    """
    Create a mapping of custom styles based on class attributes.
    """
    styles = document.styles

    # Example: Define a style for 'sc' (small caps)
    if "SmallCaps" not in styles:
        style = styles.add_style("SmallCaps", WD_STYLE_TYPE.CHARACTER)
        style.font.small_caps = True

    # Add more styles as needed based on class attributes
    # For example, styles for 's21', 's22', etc.


def process_text_run(element, paragraph, text, style_names=None):
    """
    Processes a text run, applying formatting based on the element tag and style names.
    """
    run = paragraph.add_run(text)
    if style_names:
        for style_name in style_names:
            if style_name == "bold":
                run.bold = True
            elif style_name == "italic":
                run.italic = True
            else:
                run.style = style_name


def process_node(element, paragraph, style_names=None):
    """
    Recursively processes XML nodes, handling text and child elements.
    """
    # Initialize style_names if None
    if style_names is None:
        style_names = []

    # Handle the element's text
    if element.text:
        if element.tag == "a":
            href = element.get("href")
            if href:
                add_hyperlink(paragraph, href, element.text)
            else:
                paragraph.add_run(element.text)
        elif element.tag == "br":
            paragraph.add_run().add_break()
        else:
            current_styles = style_names.copy()
            # Apply styles based on element tag or class
            if element.tag == "b":
                current_styles.append("bold")
            elif element.tag == "i":
                current_styles.append("italic")
            elif element.tag == "span":
                class_attr = element.get("class", "")
                if "sc" in class_attr:
                    current_styles.append("SmallCaps")
                # Add more class-based styles as needed
            process_text_run(
                element, paragraph, element.text, style_names=current_styles
            )

    # Handle child elements
    for child in element:
        if child.tag == "br":
            paragraph.add_run().add_break()
        elif child.tag == "note":
            # Handle footnotes
            note_text = "".join(child.itertext())
            add_footnote(paragraph, note_text)
        else:
            process_node(child, paragraph, style_names=style_names)
        # Handle tail text after the child element
        if child.tail:
            current_styles = style_names.copy()
            process_text_run(child, paragraph, child.tail, style_names=current_styles)


def process_element(element, document):
    """
    Processes XML elements and maps them to Word document elements.
    """
    if element.tag in ["p", "para"]:
        # Apply paragraph styles based on class attribute if needed
        class_attr = element.get("class", "")
        if class_attr and class_attr in document.styles:
            paragraph = document.add_paragraph(style=class_attr)
        else:
            paragraph = document.add_paragraph(style="Normal")
        process_node(element, paragraph)
    elif element.tag in ["title", "head"]:
        level = int(element.get("level", 1))  # Default to level 1 if not specified
        text = "".join(element.itertext())
        document.add_heading(text, level=level)
    elif element.tag in ["div", "div1"]:
        # Process attributes like 'title' if needed
        title = element.get("title")
        if title:
            document.add_heading(title, level=1)
        for child in element:
            process_element(child, document)
    elif element.tag == "ThML.body":
        for child in element:
            process_element(child, document)
    else:
        for child in element:
            process_element(child, document)


def convert_thml_to_docx(thml_file, docx_file):
    """
    Main function to convert a ThML file to a Word document.
    """
    tree = ET.parse(thml_file)
    root = tree.getroot()
    document = Document()

    # Create style mappings
    get_style_mapping(document)

    process_element(root, document)
    document.save(docx_file)


if __name__ == "__main__":
    import sys

    if len(sys.argv) != 3:
        print("Usage: python thml_to_docx.py input.thml output.docx")
    else:
        thml_file = sys.argv[1]
        docx_file = sys.argv[2]
        convert_thml_to_docx(thml_file, docx_file)
        print(f"Conversion complete. The Word document is saved as '{docx_file}'.")
